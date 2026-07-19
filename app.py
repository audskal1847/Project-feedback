import streamlit as st
from openai import OpenAI
import google.generativeai as genai
import requests
from bs4 import BeautifulSoup
import io
import os
import re
import pandas as pd
from openpyxl.styles import Alignment
from docx import Document
import markdown

# 1. 페이지 기본 설정
st.set_page_config(page_title="학생부 심화 탐구주제 구체화 도우미", page_icon="🏫", layout="wide")
st.title("💡 학생부 심화 탐구주제 구체화 프로그램")
st.markdown("추상적이고 두루뭉술한 탐구 주제를 2022 개정 교육과정 교과목 및 진로 목표에 맞춰 체계적인 '꼬리물기 심화 탐구'로 다듬어 드립니다.")

# --- 세션 상태 초기화 (결과 유지용) ---
if "result_text" not in st.session_state:
    st.session_state.result_text = ""

# --- 웹 크롤링 함수 정의 ---
@st.cache_data(ttl=3600)
def scrape_reference_data():
    url = "https://nojaesu.com/category/DIRECTORY/%EA%B5%90%EA%B3%BC%EC%97%B0%EA%B3%84%26%EC%A0%84%EA%B3%B5%EC%A0%81%ED%95%A9%EC%84%9C%20%EA%B8%B0%EC%82%AC%20%EB%AA%A8%EC%9D%8C"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        
        extracted_texts = []
        for tag in soup.find_all(['h2', 'h3', 'a']):
            text = tag.get_text(strip=True)
            if len(text) > 5 and text not in extracted_texts:
                extracted_texts.append(text)
                
        result = " | ".join(extracted_texts[:50])
        return result
    except Exception as e:
        return f"크롤링 실패 (AI 자체 데이터 활용 요망): {e}"
# --------------------------------------------------

# 2. 사이드바: API 제공자 및 키 설정
with st.sidebar:
    st.header("🔑 API 설정")
    
    api_provider = st.radio(
        "사용할 API 제공자를 선택하세요",
        ["Google AI Studio (공식)", "OpenRouter"]
    )
    
    if api_provider == "Google AI Studio (공식)":
        api_key = st.text_input("Google API 키를 입력하세요", type="password")
        st.info("Google AI Studio는 'gemini-3.5-flash' 모델로 고정되어 실행됩니다.")
        selected_model = "gemini-3.5-flash"
        st.markdown("[Google AI Studio 키 발급받기](https://aistudio.google.com/)")
        
    else:  
        api_key = st.text_input("OpenRouter API 키를 입력하세요", type="password")
        selected_model = st.selectbox(
            "OpenRouter 모델 선택",
            [
                "anthropic/claude-sonnet-5",
                "anthropic/claude-opus-4.8",
                "google/gemini-3.5-flash",
                "deepseek/deepseek-v4-pro",
                "qwen/qwen3.7-plus",
                "z-ai/glm-5.2",
                "openai/gpt-oss-120b:free",
                "google/gemma-4-31b-it:free"
            ]
        )
        st.markdown("[OpenRouter 키 발급받기](https://openrouter.ai/)")

# 3. 인적사항 입력
st.subheader("1. 학생 정보 입력")
col1, col2, col3 = st.columns(3)
with col1:
    school_name = st.text_input("학교 이름")
with col2:
    grade = st.selectbox("학년", ["1학년", "2학년", "3학년"])
with col3:
    student_name = st.text_input("이름")

# 4. 진로 목표 및 관심 분야 입력
st.subheader("2. 진로 목표 및 관심 분야")
career_track = st.selectbox(
    "희망하는 진로 계열을 선택하세요",
    ["공학 계열", "의약학 계열", "교육 계열", "자연과학 계열", "사회과학 계열", "경영/경제 계열", "인문/어학 계열", "예체능 계열", "기타"]
)
specific_interest = st.text_input("구체적인 희망 전공이나 관심 키워드가 있다면 적어주세요 (예: 인공지능, 신약개발, 교육공학 등)")

# 5. 기존(사전) 활동 입력 및 파일 업로드 기능
st.subheader("3. 기존 수행 활동 (사전 연계 주제)")
prior_activity = st.text_area(
    "지금까지 수행했던 활동 중, 이번 탐구와 연계하고 싶은 이전 활동 내용이나 주제를 적어주세요. (없을 경우 비워두셔도 됩니다.)",
    placeholder="예) 1학년 통합과학 시간에 '신재생 에너지의 한계'에 대해 발표했었음."
)
uploaded_file = st.file_uploader("또는 학생의 기존 활동 내용이 기록된 파일을 업로드하여 연계할 수 있습니다. (.txt, .md 지원)", type=["txt", "md"])
if uploaded_file is not None:
    st.success("📄 파일이 성공적으로 업로드되었습니다. 기존 텍스트 입력값과 함께 분석에 반영됩니다.")

# 6. 연관 교과목 선택
st.subheader("4. 탐구 주제 연관 교과목 선택")
subject_categories = [
    "국어 (화법과 언어, 독서와 작문, 문학 등)",
    "수학 (대수, 미적분, 확률과 통계, 기하 등)",
    "영어 (영어, 진로 영어, 영어 독해와 작문 등)",
    "역사 (한국사, 세계사, 동아시아 역사 기행 등)",
    "지리 (한국지리 탐구, 세계시민과 지리, 여행지리 등)",
    "일반사회 (정치, 법과 사회, 경제, 사회와 문화 등)",
    "윤리 (윤리와 사상, 현대사회와 윤리, 윤리문제 탐구 등)",
    "물리학 (물리학, 역학과 에너지, 전자기와 양자 등)",
    "화학 (화학, 물질과 에너지, 화학 반응의 세계 등)",
    "생명과학 (생명과학, 세포와 물질대사, 생물의 유전 등)",
    "지구과학 (지구과학, 지구시스템과학, 행성우주과학 등)",
    "정보/소프트웨어 (정보, 인공지능 기초, 데이터 과학 등)",
    "기술·가정/공학 (창의 공학 설계, 로봇과 공학세계 등)",
    "예술/체육/교양 (음악, 미술, 체육, 보건, 철학 등)"
]
selected_subjects = st.multiselect(
    "현재 생각 중인 탐구 주제와 가장 연관성이 깊다고 생각하는 교과목을 선택해주세요. (복수 선택 가능)",
    subject_categories,
    placeholder="연관 교과목을 선택하세요..."
)

# 7. 활용 키워드 및 교과 핵심 내용
st.subheader("5. 활용하고 싶은 키워드 혹은 교과 핵심 내용 요소")
keyword_elements = st.text_input(
    "교과서에서 배운 핵심 원리나 탐구에 꼭 포함하고 싶은 키워드를 적어주세요.",
    placeholder="예) 산화·환원 반응, 촉매, 효소 작용 원리 등"
)

# 8. 희망 탐구 주제
st.subheader("6. 희망 탐구 주제")
initial_topic = st.text_area(
    "자료조사를 통해 확정한, 혹은 고민 중인 탐구 주제를 자유롭게 적어주세요.",
    placeholder="예) 2학년 화학 시간에 배운 중화반응을 활용해 토양 오염을 해결하는 방안을 탐구하고 싶음."
)

# 9. 시스템 프롬프트
system_prompt = """
당신은 고등학교 학생부종합전형(입학사정관) 및 2022 개정 교육과정 교과 세특 전문가입니다. 
학생이 입력하거나 파일로 제출한 [사전 수행 활동], [연관 교과목], [활용 키워드/교과 핵심 요소], [희망 탐구 주제]를 유기적으로 연결하여, 과거 활동에서 한 단계 더 심화·확장되는 '꼬리물기식 탐구(종단적 연계)'가 되도록 설계해야 합니다.

[핵심 출력 조건: 개조식 작성]
모든 내용은 길고 장황한 서술형 문장을 절대 배제하십시오. 
반드시 **구체적인 교과 개념 키워드 및 전공 학술 용어**를 중심으로 한 **명료한 개조식(-, • 등의 기호 활용 및 명사형/음슴체 종결)**으로만 작성하세요. (예: ~함, ~분석함, ~을 제안함)

반드시 다음 6가지 형식과 조건에 맞춰 답변을 제공하세요:

1. [교과 연계성 및 주제 진단]
   - 학생의 '사전 활동', '희망 주제', '선택한 연관 교과목', 그리고 **'활용 키워드/핵심 요소'** 간의 유기적 연계성 평가 (개조식)
   - 2022 개정 교육과정을 바탕으로 해당 교과목의 어떤 핵심 개념과 원리가 본 탐구 주제에 적용될 수 있는지 명시적으로 분석 (개조식)
   - 학업 역량을 돋보이게 하기 위한 발전 및 보완점 제시 (개조식)

2. [탐구 출발점: 핵심 질문 및 호기심 도출]
   - 본격적인 탐구 주제로 발전하기 전, 학생이 일상이나 교과 수업 중 가질 법한 자연스럽고 날카로운 **'지적 호기심'이나 '핵심 질문' 2~3가지**를 구체적인 의문문 형태로 제시 (개조식)

3. [구체화된 탐구 주제 제안]
   - 진로, 교과목, 핵심 키워드, 이전 활동과 연계된 구체적이고 심화된 탐구 소주제를 **반드시 5가지** 제안
   - 각 주제는 문제해결, 논리적 분석, 교과 개념이 명시된 전문적인 '보고서 제목' 형태로 작성

4. [탐구 과정 가이드 (상세화)]
   - 제안된 5가지 주제 중 가장 추천하는 1가지를 골라, 다음 4단계에 따라 구체적 가이드
   - [동기 및 가설 설정]: 앞서 도출한 핵심 질문이나 교과서 개념에서 출발하여 어떤 가설을 세울 수 있는지 (개조식)
   - [핵심 탐구 질문]: 이 탐구를 관통하는 구체적인 리서치 퀘스천(Research Question) (명확한 질문형)
   - [조사 및 분석 방법]: 분석 방식, 통계자료, 교과서 단원 등 전공 용어 활용 (개조식)
   - [예상 결론 및 시사점 (구체화)]: 탐구를 통해 도출될 수 있는 구체적인 결론의 모습과, 이것이 지원 전공(진로) 분야에 던지는 학술적/실무적 시사점을 상세히 작성 (개조식)

5. [탐구의 기대 효과 및 유의사항] 
   - [탐구 진행 시 긍정적 기대 효과]: 탐구 성공 시 입학사정관에게 어필할 수 있는 학업 역량, 전공 적합성 등 구체적인 성장 포인트 (개조식)
   - [예상되는 어려움 및 극복 방안]: 자료 조사의 한계, 어려운 개념 등 겪을 수 있는 구체적인 어려움과 현실적인 팁(Tip) 안내 (개조식)

6. [맞춤형 도서 및 연계 활동 제안]
   - [추천 도서/기사]: 제공된 **[참고 웹 크롤링 데이터]**를 최우선 분석하여, 고등학생 수준에 맞는 관련 도서/기사 **2권(편)** 추천 (이유 포함)
   - [일반 후속 연계 활동]: 창체나 타 과목 세특으로 확장할 수 있는 후속 활동 **반드시 3가지** 제안 (개조식)
   - [독서 융합 심화 연계 활동]: 앞서 추천한 도서 내용과 현재 주제를 융합하는 심화 독서 연계 활동 **1가지** 제안 (개조식)
"""

# 10. 피드백 생성 버튼 및 실행 로직
if st.button("🚀 교과 연계 심화 탐구 피드백 받기"):
    if not api_key:
        st.error("왼쪽 사이드바에 해당 제공자의 API 키를 입력해주세요.")
    elif not selected_subjects:
        st.warning("4번 항목에서 연관 교과목을 최소 1개 이상 선택해주세요.")
    elif not initial_topic:
        st.warning("6번 항목에 희망 탐구 주제를 입력해주세요.")
    else:
        with st.spinner("2022 개정 교육과정 세부 분석 및 핵심 키워드를 바탕으로 심층 분석 중입니다... ⏳"):
            try:
                # 1) 웹 크롤링 수행
                crawled_data = scrape_reference_data()
                
                # 2) 파일 및 텍스트 데이터 병합 처리
                prior_text = prior_activity if prior_activity.strip() else ""
                if uploaded_file is not None:
                    try:
                        file_content = uploaded_file.read().decode("utf-8")
                        if prior_text:
                            prior_text += f"\n\n[추가 학생 제출 파일 내용]\n{file_content}"
                        else:
                            prior_text = file_content
                    except Exception as file_err:
                        st.error(f"업로드된 파일을 읽는 중 오류가 발생했습니다: {file_err}")
                        st.stop()
                
                if not prior_text.strip():
                    prior_text = "특별히 입력된 사전 활동 없음 (현재 선택된 교과목과 주제에 집중하여 심화할 것)"
                    
                subjects_str = ", ".join(selected_subjects)
                keyword_str = keyword_elements if keyword_elements.strip() else "특별히 지정된 키워드 없음"
                
                user_prompt = f"""
                [학생 정보] {school_name} {grade} {student_name}
                [진로 계열] {career_track} ({specific_interest})
                [사전 수행 활동 및 파일 본문] {prior_text}
                [선택한 연관 교과목] {subjects_str}
                [활용 키워드/교과 핵심 요소] {keyword_str}
                [희망 탐구 주제] {initial_topic}
                
                [참고 웹 크롤링 데이터]
                {crawled_data}
                """
                
                # 3) 모델 호출
                if api_provider == "Google AI Studio (공식)":
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel(
                        model_name=selected_model,
                        system_instruction=system_prompt
                    )
                    response = model.generate_content(user_prompt)
                    st.session_state.result_text = response.text
                else:
                    client = OpenAI(
                        base_url="https://openrouter.ai/api/v1",
                        api_key=api_key,
                    )
                    response = client.chat.completions.create(
                        model=selected_model, 
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt}
                        ]
                    )
                    st.session_state.result_text = response.choices[0].message.content

            except Exception as e:
                st.error(f"오류가 발생했습니다: {e}")

# 11. 결과 출력 및 안정화된 3종(Word, Excel, HTML) 다운로드 기능 제공
if st.session_state.result_text:
    st.success("분석이 완료되었습니다!")
    st.markdown("---")
    st.markdown(st.session_state.result_text)
    
    st.markdown("---")
    st.markdown("### 📥 결과 다운로드")
    
    # -----------------------------------------------------------------
    # (1) Word 파일 생성 함수 (마크다운 기호를 실제 워드 스타일로 완벽 변환)
    # -----------------------------------------------------------------
    def process_bold_text(paragraph, text):
        """ **텍스트** 를 워드의 Bold로 변환하는 함수 """
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for i, part in enumerate(parts):
            if i % 2 == 1: # ** 안의 텍스트
                paragraph.add_run(part).bold = True
            else:
                paragraph.add_run(part)

    def create_word(text):
        doc = Document()
        doc.add_heading('학생부 심화 탐구주제 구체화 결과', 0)
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Markdown 해딩 처리
            if line.startswith('### '):
                p = doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                p = doc.add_heading(line[2:].strip(), level=1)
            # Markdown 리스트 처리
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                process_bold_text(p, line[2:].strip())
            else:
                p = doc.add_paragraph()
                process_bold_text(p, line)
                
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # -----------------------------------------------------------------
    # (2) Excel 파일 생성 함수 (섹션별로 나누어 텍스트 줄바꿈 적용)
    # -----------------------------------------------------------------
    def create_excel(text):
        sections = []
        current_section = "개요"
        current_content = []
        
        for line in text.split('\n'):
            line = line.strip()
            # "1. [제목]" 또는 "## 1. [제목]" 형태 탐지하여 엑셀 행 분리
            match = re.match(r'^(?:#+ )?(\d+\.\s*\[.*?\])', line)
            if match:
                if current_content:
                    sections.append({"구분": current_section, "내용": "\n".join(current_content).strip()})
                current_section = match.group(1)
                current_content = []
            else:
                # 엑셀 지저분한 마크다운 기호 제거
                clean_line = line.replace('**', '')
                clean_line = re.sub(r'^#+\s*', '', clean_line)
                if clean_line:
                    current_content.append(clean_line)
        
        if current_content:
            sections.append({"구분": current_section, "내용": "\n".join(current_content).strip()})
            
        if not sections:
             sections = [{"구분": "탐구 결과", "내용": text.replace('**', '')}]
             
        df = pd.DataFrame(sections)
        bio = io.BytesIO()
        with pd.ExcelWriter(bio, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Result')
            worksheet = writer.sheets['Result']
            
            # 컬럼 너비 및 텍스트 줄바꿈(Wrap Text) 적용
            worksheet.column_dimensions['A'].width = 30
            worksheet.column_dimensions['B'].width = 100
            
            for row in worksheet.iter_rows(min_row=1):
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical='top')
                    
        return bio.getvalue()

    # -----------------------------------------------------------------
    # (3) HTML 파일 생성 함수 (마크다운 완벽 렌더링)
    # -----------------------------------------------------------------
    def create_html(text):
        html_body = markdown.markdown(text)
        html_content = f"""
        <!DOCTYPE html>
        <html lang="ko">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>학생부 심화 탐구주제 구체화 결과</title>
            <style>
                body {{ font-family: 'Malgun Gothic', 'Apple SD Gothic Neo', sans-serif; line-height: 1.6; padding: 20px; max-width: 800px; margin: auto; color: #333; }}
                h1 {{ color: #2c3e50; text-align: center; border-bottom: 2px solid #34495e; padding-bottom: 10px; margin-bottom: 20px; }}
                h2, h3 {{ color: #34495e; margin-top: 30px; }}
                .content {{ background-color: #f8f9fa; padding: 30px; border-radius: 8px; border: 1px solid #ececec; box-shadow: 0 4px 6px rgba(0,0,0,0.05); }}
                ul, ol {{ padding-left: 20px; }}
                li {{ margin-bottom: 8px; }}
            </style>
        </head>
        <body>
            <h1>학생부 심화 탐구주제 구체화 결과</h1>
            <div class="content">{html_body}</div>
        </body>
        </html>
        """
        return html_content.encode('utf-8')

    # 다운로드 버튼 레이아웃 배치
    dl_col1, dl_col2, dl_col3 = st.columns(3)
    
    with dl_col1:
        html_data = create_html(st.session_state.result_text)
        st.download_button(
            label="🌐 HTML (.html) 다운로드",
            data=html_data,
            file_name="탐구주제_분석결과.html",
            mime="text/html",
            use_container_width=True
        )

    with dl_col2:
        word_data = create_word(st.session_state.result_text)
        st.download_button(
            label="📄 Word (.docx) 다운로드",
            data=word_data,
            file_name="탐구주제_분석결과.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
    with dl_col3:
        excel_data = create_excel(st.session_state.result_text)
        st.download_button(
            label="📊 Excel (.xlsx) 다운로드",
            data=excel_data,
            file_name="탐구주제_분석결과.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

# 12. 프로그램 하단 만든 이 및 제목 표시 푸터
st.markdown("---")
st.markdown(
    """
    <div style="text-align: center; color: #7f8c8d; font-size: 0.9em; line-height: 1.6;">
        <strong>🏫 탐구 역량 강화를 위한 주제 탐구 구체화 어시스트 프로그램 v2.4</strong><br>
        <span style="font-size: 0.85em;">만든 이: <strong>G.E.M.S</strong></span>
    </div>
    """, 
    unsafe_allow_html=True
)
